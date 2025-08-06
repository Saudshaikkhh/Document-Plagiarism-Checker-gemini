import os
import re
import time
import random
import json
import hashlib
from datetime import datetime
from docx import Document
from PyPDF2 import PdfReader
from google.api_core import retry
import google.generativeai as genai
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
from reportlab.lib.units import inch
from docx2pdf import convert
from io import BytesIO
import tempfile

# Gemini AI setup
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel('gemini-2.5-pro')

# Cache setup
CACHE_FILE = "plagiarism_cache.json"
cache = {}

# --- Cache Management Functions ---
def load_cache():
    """Load cache from JSON file if it exists"""
    global cache
    try:
        if os.path.exists(CACHE_FILE):
            with open(CACHE_FILE, 'r') as f:
                cache = json.load(f)
                print(f"📦 Loaded cache with {len(cache)} entries")
        else:
            cache = {}
            print("🆕 Created new cache")
    except Exception as e:
        print(f"⚠️ Cache loading error: {str(e)}")
        cache = {}

def save_cache():
    """Save cache to JSON file"""
    try:
        with open(CACHE_FILE, 'w') as f:
            json.dump(cache, f, indent=2)
        print(f"💾 Saved cache with {len(cache)} entries")
    except Exception as e:
        print(f"⚠️ Cache saving error: {str(e)}")

def get_content_hash(content):
    """Compute SHA-256 hash of normalized content"""
    # Normalize content: remove extra whitespace, ignore formatting
    normalized = re.sub(r'\s+', ' ', content).strip().lower()
    return hashlib.sha256(normalized.encode('utf-8')).hexdigest()

# --- Improved A.C. Extraction Helper Functions ---
def extract_text_from_docx(docx_path):
    """Extract all text content from DOCX in document order"""
    doc = Document(docx_path)
    full_text = []
    
    # Function to process document elements in order
    def iter_block_items(parent):
        """
        Yield each paragraph and table child within *parent*, in document order.
        Each returned value is an instance of either Table or Paragraph.
        """
        from docx.document import Document as _Document
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table, _Cell
        from docx.text.paragraph import Paragraph

        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("Unsupported parent type")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)
    
    # Process all elements in order
    for block in iter_block_items(doc):
        if hasattr(block, 'text'):  # This is a Paragraph
            text = block.text.strip()
            if text:
                full_text.append(text)
                # Show first line of paragraph content
                first_line = text.split('\n')[0]
                if len(first_line) > 100:
                    first_line = first_line[:100] + "..."
                print(f"   📝 Paragraph: {first_line}")
        elif hasattr(block, 'rows'):  # This is a Table
            print(f"   📊 Table with {len(block.rows)} rows")
            for row in block.rows:
                for cell in row.cells:
                    cell_text = []
                    for para in cell.paragraphs:
                        if para.text.strip():
                            cell_text.append(para.text.strip())
                    if cell_text:
                        full_text.extend(cell_text)
                        # Show first line of table cell content
                        first_line = cell_text[0].split('\n')[0]
                        if len(first_line) > 80:
                            first_line = first_line[:80] + "..."
                        print(f"      ▸ Cell: {first_line}")
    
    return "\n".join(full_text)

def extract_text_from_pdf(pdf_source):
    """Extract all text content from PDF with detailed logging"""
    # Handle both file paths and BytesIO objects
    if isinstance(pdf_source, str):
        print(f"📄 Processing PDF file: {pdf_source}")
        reader = PdfReader(pdf_source)
    else:
        print("📄 Processing in-memory PDF (converted from DOCX)")
        pdf_source.seek(0)  # Ensure we're at the start of the stream
        reader = PdfReader(pdf_source)
    
    full_text = []
    print(f"📄 PDF has {len(reader.pages)} pages")
    
    for page_num, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text:
            full_text.append(text)
            # Show first few lines of page content
            lines = text.split('\n')
            print(f"   📄 Page {page_num}:")
            for i, line in enumerate(lines[:3]):  # Show first 3 lines of page
                if line.strip():
                    print(f"      {i+1}. {line.strip()[:80]}{'...' if len(line) > 80 else ''}")
    
    return "\n".join(full_text)

def extract_ac_sections(text_content):
    """Extract A.C. sections from text content using enhanced pattern matching"""
    print(f"📝 Total text length: {len(text_content)} characters")
    
    # Enhanced patterns to match various A.C. section formats
    ac_patterns = [
        # Traditional A.C. patterns
        re.compile(r'\bA\.?C\.?\s*(\d+\.\d+)\s+.*?[:\-]?\s*\n?', re.IGNORECASE | re.MULTILINE),
        re.compile(r'\bA\.?C\.?\s*(\d+\.\d+)\b', re.IGNORECASE),
        re.compile(r'\bAC\s*(\d+\.\d+)', re.IGNORECASE),
        re.compile(r'Assessment\s+Criteria\s+(\d+\.\d+)', re.IGNORECASE),
        
        # Direct numbering patterns (like in your document)
        re.compile(r'^\*?\*?(\d+\.\d+)\s+[A-Z][^*]*\*?\*?\s*$', re.MULTILINE),  # **1.1 Title**
        re.compile(r'^\s*(\d+\.\d+)\s+[A-Z][a-z].*?$', re.MULTILINE),  # 1.1 Title
        re.compile(r'^\s*\*\*(\d+\.\d+)\s+.*?\*\*\s*$', re.MULTILINE),  # **1.1 ...**
        re.compile(r'(?:^|\n)\s*(\d+\.\d+)\s+[A-Z][a-z].*?(?:\n|$)', re.MULTILINE),  # Flexible line matching
        
        # Table-based patterns
        re.compile(r'(\d+\.\d+)\s+COVERED[:\-]?\s*\n?', re.IGNORECASE),
        re.compile(r'(\d+\.\d+)\s*[A-Z]', re.IGNORECASE),
        re.compile(r'(\d+\.\d+)\s*:', re.IGNORECASE),
        re.compile(r'(\d+\.\d+)\s*[-–]', re.IGNORECASE),
        
        # More flexible patterns
        re.compile(r'\b(\d+\.\d+)\b.*?(?=\d+\.\d+|\Z)', re.IGNORECASE | re.DOTALL),  # Everything between numbers
    ]
    
    # Find all potential section headers with their positions
    section_candidates = []
    
    for pattern_idx, pattern in enumerate(ac_patterns):
        print(f"🔍 Testing pattern {pattern_idx + 1}: {pattern.pattern}")
        matches = list(pattern.finditer(text_content))
        print(f"   Found {len(matches)} matches")
        
        for match in matches:
            start_pos = match.start()
            ac_num = match.group(1)
            
            # Get surrounding context to validate this is a real section header
            context_start = max(0, start_pos - 100)
            context_end = min(len(text_content), start_pos + 200)
            context = text_content[context_start:context_end]
            
            # Check if this looks like a real section header
            is_valid_header = (
                # Should have substantial content after it
                len(text_content[match.end():match.end()+50].strip()) > 10 and
                # Should not be in the middle of a sentence
                (start_pos == 0 or text_content[start_pos-1] in '\n\r\t ') and
                # AC number should be reasonable (1.1 to 9.9)
                1 <= int(ac_num.split('.')[0]) <= 9 and
                1 <= int(ac_num.split('.')[1]) <= 9
            )
            
            if is_valid_header:
                section_candidates.append({
                    'start_pos': start_pos,
                    'ac_num': ac_num,
                    'pattern_idx': pattern_idx,
                    'match_text': match.group(0),
                    'context': context
                })
                print(f"   ✅ Valid candidate: A.C. {ac_num} at position {start_pos}")
                print(f"      Match: {match.group(0)[:50]}...")
    
    # Remove duplicates (same AC number found by multiple patterns)
    unique_sections = {}
    for candidate in section_candidates:
        ac_num = candidate['ac_num']
        if ac_num not in unique_sections or candidate['start_pos'] < unique_sections[ac_num]['start_pos']:
            unique_sections[ac_num] = candidate
    
    # Sort by position in document
    sorted_sections = sorted(unique_sections.values(), key=lambda x: x['start_pos'])
    
    print(f"📋 Found {len(sorted_sections)} unique A.C. sections")
    
    # Extract content for each section
    sections = {}
    for i, section in enumerate(sorted_sections):
        start_pos = section['start_pos']
        ac_num = section['ac_num']
        
        # Determine end of section (start of next section or end of document)
        if i < len(sorted_sections) - 1:
            end_pos = sorted_sections[i+1]['start_pos']
        else:
            end_pos = len(text_content)
        
        # Extract section content
        section_content = text_content[start_pos:end_pos].strip()
        
        # Clean up the content
        lines = section_content.split('\n')
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            if line and not line.startswith('==') and len(line) > 2:
                # Remove excessive formatting marks
                line = re.sub(r'\*{2,}', '', line)  # Remove multiple asterisks
                line = re.sub(r'_{2,}', '', line)  # Remove multiple underscores
                cleaned_lines.append(line)
        
        section_content = '\n'.join(cleaned_lines)
        
        # Only keep sections with substantial content
        if len(section_content) > 100:
            sections[ac_num] = section_content
            
            # Log content preview
            content_lines = section_content.split('\n')
            print(f"🔍 Extracted A.C. {ac_num} ({len(section_content)} chars):")
            for j, line in enumerate(content_lines[:3]):  # Show first 3 lines
                if line.strip():
                    print(f"   {j+1}. {line.strip()[:100]}{'...' if len(line) > 100 else ''}")
        else:
            print(f"⚠️ A.C. {ac_num} has insufficient content ({len(section_content)} chars), skipping")
    
    # If still no sections found, try aggressive fallback
    if not sections:
        print("⚠️ No sections found with enhanced patterns, trying aggressive fallback...")
        
        # Look for any numbered sections
        fallback_pattern = re.compile(r'(\d\.\d)', re.MULTILINE)
        matches = list(fallback_pattern.finditer(text_content))
        
        if matches:
            print(f"   Found {len(matches)} potential numbered sections")
            
            # Group consecutive paragraphs after each number
            for i, match in enumerate(matches):
                start_pos = match.start()
                ac_num = match.group(1)
                
                # Find end position
                if i < len(matches) - 1:
                    end_pos = matches[i+1].start()
                else:
                    end_pos = len(text_content)
                
                # Extract content
                content = text_content[start_pos:end_pos].strip()
                
                # Only keep if substantial content
                if len(content) > 200 and ac_num not in sections:
                    sections[ac_num] = content
                    print(f"   📝 Fallback found A.C. {ac_num}")
    
    if not sections:
        # Write extracted text to file for debugging
        debug_file = os.path.join(os.path.dirname(__file__), "extracted_text_debug.txt")
        with open(debug_file, "w", encoding="utf-8") as f:
            f.write(text_content)
        print(f"⚠️ No A.C. sections found! Extracted text saved to {debug_file}")
    
    return sections

def classify_text_domain(text):
    """Classify the domain/field of the text using Gemini AI"""
    
    @retry.Retry(predicate=retry.if_exception_type(Exception), deadline=60.0)
    def _make_classification_request():
        classification_prompt = f"""
        Analyze the following text and classify it into one of these academic/professional domains. 
        Return ONLY the domain name from this list:

        - Software Development
        - Web Development
        - Data Science
        - Machine Learning
        - Cybersecurity
        - Computer Networks
        - Database Management
        - Business Management
        - Marketing
        - Finance
        - Accounting
        - Human Resources
        - Healthcare
        - Education
        - Engineering
        - Science
        - Legal
        - Government
        - Non-profit
        - General Academic
        - Technical Writing
        - Creative Writing
        - Research

        Text to classify:
        {text[:2000]}

        Domain:
        """
        
        response = model.generate_content(classification_prompt)
        return response.text.strip()
    
    try:
        domain = _make_classification_request()
        print(f"🎯 Classified text domain as: {domain}")
        return domain
    except Exception as e:
        print(f"⚠️ Domain classification failed: {str(e)}")
        return "General Academic"

def check_plagiarism_gemini(content, content_hash, domain="General Academic"):
    """Check for plagiarism using Gemini AI with improved prompting and caching"""
    
    # Check cache first
    if content_hash in cache:
        print(f"💾 Cache hit for content hash: {content_hash[:8]}...")
        cached_result = cache[content_hash]
        
        # Validate cached result format
        if isinstance(cached_result, dict) and 'similarity_score' in cached_result:
            return cached_result
        else:
            print("⚠️ Invalid cached result format, regenerating...")
    
    # Enhanced prompt for plagiarism detection
    plagiarism_prompt = f"""
    You are an expert plagiarism detection system specializing in {domain} content.
    
    Analyze the following text for potential plagiarism indicators. Look for:
    
    1. **Writing Style Inconsistencies**: Sudden changes in vocabulary, sentence structure, or writing sophistication
    2. **Formatting Anomalies**: Inconsistent formatting, fonts, or spacing that might indicate copy-paste
    3. **Content Flow Issues**: Abrupt topic changes, missing context, or logical disconnections
    4. **Language Patterns**: 
       - Generic or overly polished language uncommon for student work
       - Technical terminology used inconsistently
       - Phrases that seem too advanced or specialized for the context
    5. **Structural Red Flags**: 
       - Inconsistent citation styles within the text
       - References to unnamed sources or vague attributions
       - Information that seems too detailed or specific without proper sourcing
    
    For {domain} specifically, also consider:
    - Industry-specific jargon usage patterns
    - Technical accuracy and depth appropriate for the academic level
    - Contemporary vs. outdated terminology or practices
    
    Content to analyze:
    ---
    {content}
    ---
    
    Provide your analysis in this exact JSON format:
    {{
        "similarity_score": [0-100 integer],
        "confidence_level": "[Very Low|Low|Medium|High|Very High]",
        "risk_assessment": "[Minimal|Low|Moderate|High|Critical]",
        "primary_concerns": ["list", "of", "main", "issues"],
        "detailed_analysis": "Detailed explanation of findings",
        "recommendations": "Specific recommendations for review",
        "analysis_timestamp": "{datetime.now().isoformat()}",
        "content_hash": "{content_hash}"
    }}
    """
    
    @retry.Retry(predicate=retry.if_exception_type(Exception), deadline=120.0)
    def _make_plagiarism_request():
        response = model.generate_content(plagiarism_prompt)
        return response.text.strip()
    
    try:
        response_text = _make_plagiarism_request()
        
        # Try to parse JSON response
        try:
            # Clean response text (remove markdown formatting if present)
            json_text = response_text
            if "```json" in json_text:
                json_text = json_text.split("```json")[1].split("```")[0]
            elif "```" in json_text:
                json_text = json_text.split("```")[1].split("```")[0]
            
            result = json.loads(json_text.strip())
            
            # Validate result structure
            required_fields = ['similarity_score', 'confidence_level', 'risk_assessment']
            if all(field in result for field in required_fields):
                # Cache the result
                cache[content_hash] = result
                return result
            else:
                print(f"⚠️ Invalid response structure: {result}")
                
        except json.JSONDecodeError as e:
            print(f"⚠️ JSON parsing failed: {e}")
            print(f"Raw response: {response_text[:200]}...")
            
    except Exception as e:
        print(f"⚠️ Plagiarism check failed: {str(e)}")
    
    # Fallback response
    fallback_result = {
        "similarity_score": 15,
        "confidence_level": "Medium",
        "risk_assessment": "Low",
        "primary_concerns": ["Unable to complete full analysis"],
        "detailed_analysis": "Analysis could not be completed due to technical issues. Manual review recommended.",
        "recommendations": "Please review this content manually for potential plagiarism indicators.",
        "analysis_timestamp": datetime.now().isoformat(),
        "content_hash": content_hash
    }
    
    # Cache fallback result
    cache[content_hash] = fallback_result
    return fallback_result

def detect_document_topic(content_sample):
    """Detect the main topic/domain of the document"""
    try:
        prompt = (
            "Identify the main academic or professional topic of the following document excerpt. "
            "Respond with only the topic name in 3-5 words.\n\n"
            f"EXCERPT:\n{content_sample[:2000]}"
        )
        
        response = model.generate_content(prompt)
        topic = response.text.strip()
        topic = re.sub(r'[^a-zA-Z0-9\s]', '', topic)
        return topic
    except Exception as e:
        print(f"❌ Topic detection error: {str(e)}")
        return "Academic_Subject"

def generate_pdf_report(results, document_topic, output_buffer):
    """Generate PDF report and write to buffer"""
    doc = SimpleDocTemplate(
        output_buffer,
        pagesize=letter,
        rightMargin=0.5*inch,
        leftMargin=0.5*inch,
        topMargin=0.5*inch,
        bottomMargin=0.5*inch
    )
    
    styles = getSampleStyleSheet()
    elements = []
    
    # Title
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=20,
        alignment=1,
        spaceAfter=0.3*inch,
        textColor=colors.HexColor("#2E5984")
    )
    
    elements.append(Paragraph(f"Plagiarism Detection Report - {document_topic}", title_style))
    elements.append(Spacer(1, 0.2*inch))
    
    # Summary statistics
    if results:
        total_sections = len(results)
        high_risk = sum(1 for r in results.values() if r['similarity_score'] >= 61)
        medium_risk = sum(1 for r in results.values() if 41 <= r['similarity_score'] < 61)
        low_risk = sum(1 for r in results.values() if r['similarity_score'] < 41)
        
        summary_data = [
            ['Total Sections', str(total_sections)],
            ['High Risk (≥61%)', str(high_risk)],
            ['Medium Risk (41-60%)', str(medium_risk)],
            ['Low Risk (<41%)', str(low_risk)]
        ]
        
        summary_table = Table(summary_data, colWidths=[2*inch, 1*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.grey),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,0), 12),
            ('BOTTOMPADDING', (0,0), (-1,0), 12),
            ('BACKGROUND', (0,1), (-1,-1), colors.beige),
            ('GRID', (0,0), (-1,-1), 1, colors.black)
        ]))
        
        elements.append(summary_table)
        elements.append(Spacer(1, 0.3*inch))
        
        # Detailed results table
        table_data = [['A.C. Section', 'Similarity Score', 'Risk Level', 'Primary Concerns']]
        
        for ac_num, result in sorted(results.items()):
            risk_level = "High" if result['similarity_score'] >= 61 else "Medium" if result['similarity_score'] >= 41 else "Low"
            concerns = ', '.join(result.get('primary_concerns', ['None'])[:3])
            
            table_data.append([
                ac_num,
                f"{result['similarity_score']}%",
                risk_level,
                concerns[:50] + "..." if len(concerns) > 50 else concerns
            ])
        
        results_table = Table(table_data, colWidths=[1*inch, 1*inch, 1*inch, 3.5*inch])
        results_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.grey),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,0), 10),
            ('BOTTOMPADDING', (0,0), (-1,0), 12),
            ('BACKGROUND', (0,1), (-1,-1), colors.beige),
            ('GRID', (0,0), (-1,-1), 1, colors.black),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('FONTSIZE', (0,1), (-1,-1), 9)
        ]))
        
        elements.append(results_table)
    
    doc.build(elements)
    return output_buffer

def main_plagiarism_pipeline(uploaded_file, progress_callback=None):
    """Main pipeline for processing uploaded files and generating plagiarism reports"""
    
    def update_progress(message, progress):
        if progress_callback:
            progress_callback(message, progress)
        print(f"📊 {message} ({progress}%)")
    
    try:
        # Initialize
        update_progress("Initializing analysis system", 5)
        load_cache()
        
        start_time = time.time()
        cache_hits = 0
        
        # Save uploaded file temporarily
        update_progress("Processing uploaded file", 10)
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as temp_file:
            temp_file.write(uploaded_file.getvalue())
            temp_file_path = temp_file.name
        
        try:
            # Extract text and A.C. sections
            update_progress("Extracting document content", 20)
            
            file_extension = uploaded_file.name.split('.')[-1].lower()
            
            if file_extension == "docx":
                # Convert DOCX to PDF in memory and extract
                with tempfile.TemporaryDirectory() as temp_dir:
                    temp_pdf_path = os.path.join(temp_dir, "temp.pdf")
                    convert(temp_file_path, temp_pdf_path)
                    
                    with open(temp_pdf_path, 'rb') as pdf_file:
                        pdf_bytes = pdf_file.read()
                
                pdf_stream = BytesIO(pdf_bytes)
                full_text = extract_text_from_pdf(pdf_stream)
                
            elif file_extension == "pdf":
                full_text = extract_text_from_pdf(temp_file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            update_progress("Identifying A.C. sections", 30)
            ac_sections = extract_ac_sections(full_text)
            
            if not ac_sections:
                raise ValueError("No A.C. sections found in the document")
            
            # Classify document domain
            update_progress("Classifying document domain", 35)
            sample_content = next(iter(ac_sections.values()))
            domain = classify_text_domain(sample_content)
            
            # Process each A.C. section
            plagiarism_results = {}
            total_sections = len(ac_sections)
            
            for i, (ac_num, content) in enumerate(ac_sections.items()):
                section_progress = 40 + (i * 40 // total_sections)
                update_progress(f"Analyzing A.C. {ac_num}", section_progress)
                
                content_hash = get_content_hash(content)
                
                # Check if cached
                if content_hash in cache:
                    cache_hits += 1
                
                result = check_plagiarism_gemini(content, content_hash, domain)
                plagiarism_results[ac_num] = result
            
            # Generate PDF report
            update_progress("Generating PDF report", 85)
            
            report_buffer = BytesIO()
            generate_pdf_report(plagiarism_results, domain, report_buffer)
            report_buffer.seek(0)
            pdf_data = report_buffer.getvalue()
            
            # Calculate final statistics
            end_time = time.time()
            total_time = end_time - start_time
            
            update_progress("Analysis complete", 100)
            
            # Save cache
            save_cache()
            
            return {
                'success': True,
                'ac_sections': ac_sections,
                'plagiarism_results': plagiarism_results,
                'domain': domain,
                'pdf_report': pdf_data,
                'processing_stats': {
                    'total_time': total_time,
                    'cache_hits': cache_hits,
                    'sections_processed': len(plagiarism_results)
                }
            }
            
        finally:
            # Clean up temporary file
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
                
    except Exception as e:
        error_msg = f"Analysis failed: {str(e)}"
        update_progress(error_msg, 0)
        return {
            'success': False,
            'error': error_msg
        }

# --- Improved A.C. Extraction Helper Functions ---
def extract_text_from_docx(docx_path):
    """Extract all text content from DOCX in document order"""
    doc = Document(docx_path)
    full_text = []
    
    # Function to process document elements in order
    def iter_block_items(parent):
        """
        Yield each paragraph and table child within *parent*, in document order.
        Each returned value is an instance of either Table or Paragraph.
        """
        from docx.document import Document as _Document
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table, _Cell
        from docx.text.paragraph import Paragraph

        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("Unsupported parent type")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)
    
    # Process all elements in order
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                full_text.append(text)
                # Show first line of paragraph content
                first_line = text.split('\n')[0]
                if len(first_line) > 100:
                    first_line = first_line[:100] + "..."
                print(f"   📝 Paragraph: {first_line}")
        elif isinstance(block, Table):
            print(f"   📊 Table with {len(block.rows)} rows")
            for row in block.rows:
                for cell in row.cells:
                    cell_text = []
                    for para in cell.paragraphs:
                        if para.text.strip():
                            cell_text.append(para.text.strip())
                    if cell_text:
                        full_text.extend(cell_text)
                        # Show first line of table cell content
                        first_line = cell_text[0].split('\n')[0]
                        if len(first_line) > 80:
                            first_line = first_line[:80] + "..."
                        print(f"      ▸ Cell: {first_line}")
    
    return "\n".join(full_text)

def extract_text_from_pdf(pdf_source):
    """Extract all text content from PDF with detailed logging"""
    # Handle both file paths and BytesIO objects
    if isinstance(pdf_source, str):
        print(f"📄 Processing PDF file: {pdf_source}")
        reader = PdfReader(pdf_source)
    else:
        print("📄 Processing in-memory PDF (converted from DOCX)")
        pdf_source.seek(0)  # Ensure we're at the start of the stream
        reader = PdfReader(pdf_source)
    
    full_text = []
    print(f"📄 PDF has {len(reader.pages)} pages")
    
    for page_num, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text:
            full_text.append(text)
            # Show first few lines of page content
            lines = text.split('\n')
            print(f"   📄 Page {page_num}:")
            for i, line in enumerate(lines[:3]):  # Show first 3 lines of page
                if line.strip():
                    print(f"      {i+1}. {line.strip()[:80]}{'...' if len(line) > 80 else ''}")
    
    return "\n".join(full_text)

def extract_ac_sections(text_content):
    """Extract A.C. sections from text content using enhanced pattern matching"""
    print(f"📝 Total text length: {len(text_content)} characters")
    
    # Enhanced patterns to match various A.C. section formats
    ac_patterns = [
        # Traditional A.C. patterns
        re.compile(r'\bA\.?C\.?\s*(\d+\.\d+)\s+.*?[:\-]?\s*\n?', re.IGNORECASE | re.MULTILINE),
        re.compile(r'\bA\.?C\.?\s*(\d+\.\d+)\b', re.IGNORECASE),
        re.compile(r'\bAC\s*(\d+\.\d+)', re.IGNORECASE),
        re.compile(r'Assessment\s+Criteria\s+(\d+\.\d+)', re.IGNORECASE),
        
        # Direct numbering patterns (like in your document)
        re.compile(r'^\*?\*?(\d+\.\d+)\s+[A-Z][^*]*\*?\*?\s*$', re.MULTILINE),  # **1.1 Title**
        re.compile(r'^\s*(\d+\.\d+)\s+[A-Z][a-z].*?$', re.MULTILINE),  # 1.1 Title
        re.compile(r'^\s*\*\*(\d+\.\d+)\s+.*?\*\*\s*$', re.MULTILINE),  # **1.1 ...**
        re.compile(r'(?:^|\n)\s*(\d+\.\d+)\s+[A-Z][a-z].*?(?:\n|$)', re.MULTILINE),  # Flexible line matching
        
        # Table-based patterns
        re.compile(r'(\d+\.\d+)\s+COVERED[:\-]?\s*\n?', re.IGNORECASE),
        re.compile(r'(\d+\.\d+)\s*[A-Z]', re.IGNORECASE),
        re.compile(r'(\d+\.\d+)\s*:', re.IGNORECASE),
        re.compile(r'(\d+\.\d+)\s*[-–]', re.IGNORECASE),
        
        # More flexible patterns
        re.compile(r'\b(\d+\.\d+)\b.*?(?=\d+\.\d+|\Z)', re.IGNORECASE | re.DOTALL),  # Everything between numbers
    ]
    
    # Find all potential section headers with their positions
    section_candidates = []
    
    for pattern_idx, pattern in enumerate(ac_patterns):
        print(f"🔍 Testing pattern {pattern_idx + 1}: {pattern.pattern}")
        matches = list(pattern.finditer(text_content))
        print(f"   Found {len(matches)} matches")
        
        for match in matches:
            start_pos = match.start()
            ac_num = match.group(1)
            
            # Get surrounding context to validate this is a real section header
            context_start = max(0, start_pos - 100)
            context_end = min(len(text_content), start_pos + 200)
            context = text_content[context_start:context_end]
            
            # Check if this looks like a real section header
            is_valid_header = (
                # Should have substantial content after it
                len(text_content[match.end():match.end()+50].strip()) > 10 and
                # Should not be in the middle of a sentence
                (start_pos == 0 or text_content[start_pos-1] in '\n\r\t ') and
                # AC number should be reasonable (1.1 to 9.9)
                1 <= int(ac_num.split('.')[0]) <= 9 and
                1 <= int(ac_num.split('.')[1]) <= 9
            )
            
            if is_valid_header:
                section_candidates.append({
                    'start_pos': start_pos,
                    'ac_num': ac_num,
                    'pattern_idx': pattern_idx,
                    'match_text': match.group(0),
                    'context': context
                })
                print(f"   ✅ Valid candidate: A.C. {ac_num} at position {start_pos}")
                print(f"      Match: {match.group(0)[:50]}...")
    
    # Remove duplicates (same AC number found by multiple patterns)
    unique_sections = {}
    for candidate in section_candidates:
        ac_num = candidate['ac_num']
        if ac_num not in unique_sections or candidate['start_pos'] < unique_sections[ac_num]['start_pos']:
            unique_sections[ac_num] = candidate
    
    # Sort by position in document
    sorted_sections = sorted(unique_sections.values(), key=lambda x: x['start_pos'])
    
    print(f"📋 Found {len(sorted_sections)} unique A.C. sections")
    
    # Extract content for each section
    sections = {}
    for i, section in enumerate(sorted_sections):
        start_pos = section['start_pos']
        ac_num = section['ac_num']
        
        # Determine end of section (start of next section or end of document)
        if i < len(sorted_sections) - 1:
            end_pos = sorted_sections[i+1]['start_pos']
        else:
            end_pos = len(text_content)
        
        # Extract section content
        section_content = text_content[start_pos:end_pos].strip()
        
        # Clean up the content
        lines = section_content.split('\n')
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            if line and not line.startswith('==') and len(line) > 2:
                # Remove excessive formatting marks
                line = re.sub(r'\*{2,}', '', line)  # Remove multiple asterisks
                line = re.sub(r'_{2,}', '', line)  # Remove multiple underscores
                cleaned_lines.append(line)
        
        section_content = '\n'.join(cleaned_lines)
        
        # Only keep sections with substantial content
        if len(section_content) > 100:
            sections[ac_num] = section_content
            
            # Log content preview
            content_lines = section_content.split('\n')
            print(f"🔍 Extracted A.C. {ac_num} ({len(section_content)} chars):")
            for j, line in enumerate(content_lines[:3]):  # Show first 3 lines
                if line.strip():
                    print(f"   {j+1}. {line.strip()[:100]}{'...' if len(line) > 100 else ''}")
        else:
            print(f"⚠️ A.C. {ac_num} has insufficient content ({len(section_content)} chars), skipping")
    
    # If still no sections found, try aggressive fallback
    if not sections:
        print("⚠️ No sections found with enhanced patterns, trying aggressive fallback...")
        
        # Look for any numbered sections
        fallback_pattern = re.compile(r'(\d\.\d)', re.MULTILINE)
        matches = list(fallback_pattern.finditer(text_content))
        
        if matches:
            print(f"   Found {len(matches)} potential numbered sections")
            
            # Group consecutive paragraphs after each number
            for i, match in enumerate(matches):
                start_pos = match.start()
                ac_num = match.group(1)
                
                # Find end position
                if i < len(matches) - 1:
                    end_pos = matches[i+1].start()
                else:
                    end_pos = len(text_content)
                
                # Extract content
                content = text_content[start_pos:end_pos].strip()
                
                # Only keep if substantial content
                if len(content) > 200 and ac_num not in sections:
                    sections[ac_num] = content
                    print(f"   📝 Fallback found A.C. {ac_num}")
    
    if not sections:
        # Write extracted text to file for debugging
        debug_file = os.path.join(os.path.dirname(__file__), "extracted_text_debug.txt")
        with open(debug_file, "w", encoding="utf-8") as f:
            f.write(text_content)
        print(f"⚠️ No A.C. sections found! Extracted text saved to {debug_file}")
    
    return sections

def classify_text_domain(text):
    """Classify the domain/field of the text using Gemini AI"""
    
    @retry.Retry(predicate=retry.if_exception_type(Exception), deadline=60.0)
    def _make_classification_request():
        classification_prompt = f"""
        Analyze the following text and classify it into one of these academic/professional domains. 
        Return ONLY the domain name from this list:

        - Software Development
        - Web Development
        - Data Science
        - Machine Learning
        - Cybersecurity
        - Computer Networks
        - Database Management
        - Business Management
        - Marketing
        - Finance
        - Accounting
        - Human Resources
        - Healthcare
        - Education
        - Engineering
        - Science
        - Legal
        - Government
        - Non-profit
        - General Academic
        - Technical Writing
        - Creative Writing
        - Research

        Text to classify:
        {text[:2000]}

        Domain:
        """
        
        response = model.generate_content(classification_prompt)
        return response.text.strip()
    
    try:
        domain = _make_classification_request()
        print(f"🎯 Classified text domain as: {domain}")
        return domain
    except Exception as e:
        print(f"⚠️ Domain classification failed: {str(e)}")
        return "General Academic"

def check_plagiarism_gemini(content, content_hash, domain="General Academic"):
    """Check for plagiarism using Gemini AI with improved prompting and caching"""
    
    # Check cache first
    if content_hash in cache:
        print(f"💾 Cache hit for content hash: {content_hash[:8]}...")
        cached_result = cache[content_hash]
        
        # Validate cached result format
        if isinstance(cached_result, dict) and 'similarity_score' in cached_result:
            return cached_result
        else:
            print("⚠️ Invalid cached result format, regenerating...")
    
    # Enhanced prompt for plagiarism detection
    plagiarism_prompt = f"""
    You are an expert plagiarism detection system specializing in {domain} content.
    
    Analyze the following text for potential plagiarism indicators. Look for:
    
    1. **Writing Style Inconsistencies**: Sudden changes in vocabulary, sentence structure, or writing sophistication
    2. **Formatting Anomalies**: Inconsistent formatting, fonts, or spacing that might indicate copy-paste
    3. **Content Flow Issues**: Abrupt topic changes, missing context, or logical disconnections
    4. **Language Patterns**: 
       - Generic or overly polished language uncommon for student work
       - Technical terminology used inconsistently
       - Phrases that seem too advanced or specialized for the context
    5. **Structural Red Flags**: 
       - Inconsistent citation styles within the text
       - References to unnamed sources or vague attributions
       - Information that seems too detailed or specific without proper sourcing
    
    For {domain} specifically, also consider:
    - Industry-specific jargon usage patterns
    - Technical accuracy and depth appropriate for the academic level
    - Contemporary vs. outdated terminology or practices
    
    Content to analyze:
    ---
    {content}
    ---
    
    Provide your analysis in this exact JSON format:
    {{
        "similarity_score": [0-100 integer],
        "confidence_level": "[Very Low|Low|Medium|High|Very High]",
        "primary_concerns": [
            "Brief description of main concern 1",
            "Brief description of main concern 2",
            "Brief description of main concern 3"
        ],
        "detailed_analysis": "Comprehensive explanation of findings including specific examples from the text",
        "recommendations": "Specific actions or areas to investigate further",
        "risk_assessment": "[Low Risk|Medium Risk|High Risk|Critical Risk]"
    }}
    
    Scoring Guidelines:
    - 0-20: No significant plagiarism indicators
    - 21-40: Minor concerns, possibly coincidental
    - 41-60: Moderate concerns requiring review
    - 61-80: High probability of plagiarism
    - 81-100: Very high probability of plagiarism
    """
    
    @retry.Retry(predicate=retry.if_exception_type(Exception), deadline=120.0)
    def _make_plagiarism_request():
        response = model.generate_content(plagiarism_prompt)
        return response.text
    
    try:
        print(f"🤖 Analyzing content with Gemini AI (Domain: {domain})...")
        
        # Add some delay to respect rate limits
        time.sleep(random.uniform(1, 3))
        
        response_text = _make_plagiarism_request()
        print(f"✅ Received response from Gemini AI")
        
        # Try to extract JSON from response
        try:
            # Look for JSON block in the response
            json_match = re.search(r'```json\s*(.*?)\s*```', response_text, re.DOTALL)
            if json_match:
                json_text = json_match.group(1)
            else:
                # Try to find JSON without code blocks
                json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
                if json_match:
                    json_text = json_match.group(0)
                else:
                    raise ValueError("No JSON found in response")
            
            result = json.loads(json_text)
            
            # Validate required fields
            required_fields = ['similarity_score', 'confidence_level', 'primary_concerns', 'detailed_analysis']
            for field in required_fields:
                if field not in result:
                    raise ValueError(f"Missing required field: {field}")
            
            # Ensure similarity_score is numeric
            result['similarity_score'] = int(result['similarity_score'])
            
            # Add metadata
            result['analysis_timestamp'] = datetime.now().isoformat()
            result['content_hash'] = content_hash
            result['domain'] = domain
            
            # Cache the result
            cache[content_hash] = result
            save_cache()
            
            print(f"📊 Plagiarism Score: {result['similarity_score']}% ({result['confidence_level']} confidence)")
            
            return result
            
        except (json.JSONDecodeError, ValueError) as e:
            print(f"⚠️ Failed to parse Gemini response as JSON: {str(e)}")
            print(f"Raw response: {response_text[:500]}...")
            
            # Create fallback result
            fallback_result = {
                'similarity_score': 50,
                'confidence_level': 'Low',
                'primary_concerns': ['Unable to analyze - API response parsing failed'],
                'detailed_analysis': f'Gemini AI response could not be parsed. Raw response: {response_text}',
                'recommendations': 'Manual review recommended due to analysis failure',
                'risk_assessment': 'Medium Risk',
                'analysis_timestamp': datetime.now().isoformat(),
                'content_hash': content_hash,
                'domain': domain
            }
            
            return fallback_result
            
    except Exception as e:
        print(f"❌ Gemini AI request failed: {str(e)}")
        
        # Create error result
        error_result = {
            'similarity_score': 0,
            'confidence_level': 'Very Low',
            'primary_concerns': [f'Analysis failed: {str(e)}'],
            'detailed_analysis': f'Could not complete plagiarism analysis due to error: {str(e)}',
            'recommendations': 'Retry analysis or use alternative detection method',
            'risk_assessment': 'Low Risk',
            'analysis_timestamp': datetime.now().isoformat(),
            'content_hash': content_hash,
            'domain': domain
        }
        
        return error_result

def generate_plagiarism_report(file_name, ac_sections, plagiarism_results, processing_stats):
    """Generate a comprehensive PDF report with improved layout and content"""
    
    # Create temporary file for the report
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    temp_file.close()
    
    try:
        doc = SimpleDocTemplate(
            temp_file.name,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Container for the report content
        story = []
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1,  # Center alignment
            textColor=colors.darkblue
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=12,
            textColor=colors.darkblue
        )
        
        subheading_style = ParagraphStyle(
            'CustomSubHeading',
            parent=styles['Heading3'],
            fontSize=12,
            spaceAfter=8,
            spaceBefore=8,
            textColor=colors.darkgreen
        )
        
        normal_style = styles['Normal']
        normal_style.fontSize = 10
        normal_style.spaceAfter = 6
        
        # Report Title
        story.append(Paragraph("Plagiarism Detection Report", title_style))
        story.append(Spacer(1, 20))
        
        # Document Information
        story.append(Paragraph("Document Information", heading_style))
        
        doc_info_data = [
            ['Document Name:', file_name],
            ['Analysis Date:', datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
            ['Total A.C. Sections:', str(len(ac_sections))],
            ['Sections Analyzed:', str(len(plagiarism_results))],
            ['Processing Time:', f"{processing_stats.get('total_time', 0):.1f} seconds"],
            ['Cache Hits:', str(processing_stats.get('cache_hits', 0))]
        ]
        
        doc_info_table = Table(doc_info_data, colWidths=[2*inch, 4*inch])
        doc_info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(doc_info_table)
        story.append(Spacer(1, 20))
        
        # Executive Summary
        story.append(Paragraph("Executive Summary", heading_style))
        
        # Calculate summary statistics
        total_sections = len(plagiarism_results)
        if total_sections > 0:
            scores = [result['similarity_score'] for result in plagiarism_results.values()]
            avg_score = sum(scores) / len(scores)
            max_score = max(scores)
            min_score = min(scores)
            
            high_risk_count = sum(1 for score in scores if score >= 61)
            medium_risk_count = sum(1 for score in scores if 41 <= score < 61)
            low_risk_count = sum(1 for score in scores if score < 41)
            
            summary_text = f"""
            This report analyzes {total_sections} A.C. sections for potential plagiarism indicators.
            
            Overall Statistics:
            • Average Similarity Score: {avg_score:.1f}%
            • Highest Score: {max_score}%
            • Lowest Score: {min_score}%
            
            Risk Distribution:
            • High Risk (61-100%): {high_risk_count} sections
            • Medium Risk (41-60%): {medium_risk_count} sections  
            • Low Risk (0-40%): {low_risk_count} sections
            """
            
            story.append(Paragraph(summary_text, normal_style))
        else:
            story.append(Paragraph("No A.C. sections were found for analysis.", normal_style))
        
        story.append(Spacer(1, 20))
        
        # Detailed Analysis for each A.C. section
        if plagiarism_results:
            story.append(Paragraph("Detailed Section Analysis", heading_style))
            
            # Sort sections by A.C. number
            sorted_acs = sorted(plagiarism_results.keys(), key=lambda x: [int(i) for i in x.split('.')])
            
            for ac_num in sorted_acs:
                result = plagiarism_results[ac_num]
                
                # Section header
                story.append(Paragraph(f"A.C. {ac_num}", subheading_style))
                
                # Create section analysis table
                section_data = [
                    ['Similarity Score:', f"{result['similarity_score']}%"],
                    ['Confidence Level:', result['confidence_level']],
                    ['Risk Assessment:', result.get('risk_assessment', 'Not specified')]
                ]
                
                section_table = Table(section_data, colWidths=[1.5*inch, 4.5*inch])
                section_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, -1), colors.lightblue),
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                story.append(section_table)
                story.append(Spacer(1, 12))
                
                # Primary Concerns
                if result.get('primary_concerns'):
                    story.append(Paragraph("<b>Primary Concerns:</b>", normal_style))
                    for concern in result['primary_concerns']:
                        story.append(Paragraph(f"• {concern}", normal_style))
                    story.append(Spacer(1, 8))
                
                # Detailed Analysis
                if result.get('detailed_analysis'):
                    story.append(Paragraph("<b>Detailed Analysis:</b>", normal_style))
                    story.append(Paragraph(result['detailed_analysis'], normal_style))
                    story.append(Spacer(1, 8))
                
                # Recommendations
                if result.get('recommendations'):
                    story.append(Paragraph("<b>Recommendations:</b>", normal_style))
                    story.append(Paragraph(result['recommendations'], normal_style))
                
                story.append(Spacer(1, 20))
                
                # Add page break for long reports
                if len(sorted_acs) > 3 and ac_num != sorted_acs[-1]:
                    story.append(PageBreak())
        
        # Report Footer
        story.append(Spacer(1, 30))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.grey,
            alignment=1
        )
        story.append(Paragraph(
            "This report was generated using AI-powered plagiarism detection. "
            "Results should be reviewed by qualified personnel and used as guidance only.",
            footer_style
        ))
        
        # Build PDF
        doc.build(story)
        
        # Read the generated PDF
        with open(temp_file.name, 'rb') as f:
            pdf_content = f.read()
        
        print(f"📄 Generated PDF report: {len(pdf_content)} bytes")
        return pdf_content
        
    except Exception as e:
        print(f"❌ PDF generation failed: {str(e)}")
        raise
    finally:
        # Clean up temporary file
        try:
            os.unlink(temp_file.name)
        except:
            pass

def main_plagiarism_pipeline(uploaded_file, progress_callback=None):
    """Main pipeline for plagiarism detection with progress tracking"""
    start_time = time.time()
    
    if progress_callback:
        progress_callback("Initializing cache...", 0)
    
    # Load cache
    load_cache()
    
    try:
        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as temp_file:
            temp_file.write(uploaded_file.getvalue())
            temp_file_path = temp_file.name
        
        if progress_callback:
            progress_callback("Extracting text from document...", 10)
        
        # Extract text based on file type
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        
        if file_extension == '.docx':
            print(f"📄 Processing DOCX file: {uploaded_file.name}")
            text_content = extract_text_from_docx(temp_file_path)
        elif file_extension == '.pdf':
            print(f"📄 Processing PDF file: {uploaded_file.name}")
            text_content = extract_text_from_pdf(temp_file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        if progress_callback:
            progress_callback("Identifying A.C. sections...", 20)
        
        # Extract A.C. sections
        ac_sections = extract_ac_sections(text_content)
        
        if not ac_sections:
            raise ValueError("No A.C. sections found in the document")
        
        if progress_callback:
            progress_callback("Classifying document domain...", 30)
        
        # Classify domain
        domain = classify_text_domain(text_content[:3000])
        
        # Analyze each section for plagiarism
        plagiarism_results = {}
        cache_hits = 0
        total_sections = len(ac_sections)
        
        for i, (ac_num, content) in enumerate(ac_sections.items()):
            progress = 30 + (i / total_sections) * 50
            if progress_callback:
                progress_callback(f"Analyzing A.C. {ac_num}...", progress)
            
            content_hash = get_content_hash(content)
            
            # Check if already cached
            if content_hash in cache:
                cache_hits += 1
            
            result = check_plagiarism_gemini(content, content_hash, domain)
            plagiarism_results[ac_num] = result
        
        if progress_callback:
            progress_callback("Generating report...", 90)
        
        # Generate report
        processing_stats = {
            'total_time': time.time() - start_time,
            'cache_hits': cache_hits,
            'total_sections': total_sections
        }
        
        pdf_content = generate_plagiarism_report(
            uploaded_file.name,
            ac_sections,
            plagiarism_results,
            processing_stats
        )
        
        if progress_callback:
            progress_callback("Analysis complete!", 100)
        
        return {
            'success': True,
            'ac_sections': ac_sections,
            'plagiarism_results': plagiarism_results,
            'processing_stats': processing_stats,
            'pdf_report': pdf_content,
            'domain': domain
        }
        
    except Exception as e:
        error_msg = f"Analysis failed: {str(e)}"
        print(f"❌ {error_msg}")
        if progress_callback:
            progress_callback(error_msg, 0)
        
        return {
            'success': False,
            'error': error_msg
        }
    
    finally:
        # Clean up temporary file
        try:
            os.unlink(temp_file_path)
        except:
            pass
